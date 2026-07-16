"""
Step 0: Parse + Render — DOCX bytes → clean markdown text.
Location: app/services/docx_v3/parser_render.py

Handles all DOCX structural complexities:
  - Nested tables (tables inside cells) → rendered as indented blocks below parent row
  - Merged cells (gridSpan / vMerge) → deduplicated
  - Inline images → [IMAGE] placeholders
  - Textboxes in drawing shapes → text extracted
  - Body elements in document order (paragraphs + tables interleaved)

Design: Nested tables are NOT crammed into parent cells (unreadable).
Instead, they are rendered as indented sub-tables below the parent row:

  | סעיף ב': ... פירוט: | 15 |
    [NESTED TABLE: 4x2]
    | כותרת הפעולה | 1 |
    |---|---|
    | מציאת הערוץ... | 3 |

This gives the LLM a clear, readable view of hierarchical rubric structures.
"""
from __future__ import annotations

import io
import logging
from collections import Counter
from dataclasses import dataclass
from typing import List, Optional, Tuple

from docx import Document
from docx.table import Table as DocxTableObj

logger = logging.getLogger(__name__)

# OOXML namespaces
_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
}


def _tag(el) -> str:
    t = el.tag
    return t.split("}")[-1] if "}" in t else t


# =============================================================================
# COLOR-CONTRAST HELPERS
# =============================================================================
# Teacher annotations (example solutions, scoring notes) are marked by COLOR. The
# load-bearing signal is LOCAL CONTRAST — "colored differently from the enclosing
# scope's dominant color" — not "is red": a uniformly-red table (e.g. csharp Q1) is
# cosmetic and must stay unmarked, while a red paragraph against black prose (899371
# solutions) must be marked. So the baseline is per-scope: table content compares to
# the table's dominant color; body prose to the body's dominant color.

# Colors at/near black are the default ink — never treated as contrast.
_BLACKISH = {"000000", "010101", "0D0D0D", "1A1A1A", "212121", "262626"}


def _is_annotation_color(hex_color: str) -> bool:
    """True only for SATURATED RED — the teacher-annotation convention in Bagrut rubrics.

    This deliberately excludes the colors DOCX uses for non-annotation purposes, which
    would otherwise flood the output: the Visual Studio C# syntax palette (008000 green
    comments, 0000FF blue keywords, 2B91AF teal types, A31515 dark-red strings) and white
    (FFFFFF) header styling. Pure bright red is disjoint from that palette, so the teacher's
    red answer-keys/notes are captured while code highlighting and table chrome are not.
    The A31515 code-string red is excluded by the R>=0xB0 floor (its R is 0xA3).
    Tunable if a teacher's convention differs.
    """
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
    except (ValueError, IndexError):
        return False
    return r >= 0xB0 and g <= 0x60 and b <= 0x60


def _run_color(r_el) -> Optional[str]:
    """Hex color of a <w:r> run from <w:rPr><w:color w:val=.../>, or None for default/auto."""
    ns_w = _NS["w"]
    rPr = r_el.find(f"{{{ns_w}}}rPr")
    if rPr is None:
        return None
    color_el = rPr.find(f"{{{ns_w}}}color")
    if color_el is None:
        return None
    val = color_el.get(f"{{{ns_w}}}val")
    if not val or val.lower() == "auto":
        return None
    return val.upper()


def _norm_color(c: Optional[str]) -> Optional[str]:
    """Collapse black-ish colors to None so they never count as contrast."""
    if c is None:
        return None
    return None if c.upper() in _BLACKISH else c.upper()


def _run_visible_len(r_el) -> int:
    ns_w = _NS["w"]
    return sum(len(t.text or "") for t in r_el.iter(f"{{{ns_w}}}t"))


# =============================================================================
# HIGHLIGHT-CONTRAST HELPERS (B1)
# =============================================================================
# Teacher annotations arrive on a SECOND physical channel: run-level HIGHLIGHT
# (<w:highlight w:val="yellow"/>), e.g. foundations q1.ג marks the correct option
# in an options list with yellow highlight and NO font color. The renderer was
# highlight-blind, so that ink never reached the LLM (and text-based render-loss
# attribution scored 0 — formatting loss is invisible to a text diff).
# Same architecture as color: LOCAL CONTRAST against the enclosing scope's
# dominant highlight, so a uniformly-highlighted table stays unmarked.

# 'white' is a paste artifact (invisible on paper — csharp has 362 such runs);
# 'none' is the explicit no-highlight value. Neither is teacher ink.
_INVISIBLE_HIGHLIGHTS = {"none", "white"}


def _run_highlight(r_el) -> Optional[str]:
    """Named highlight of a <w:r> from <w:rPr><w:highlight w:val=.../>, or None."""
    ns_w = _NS["w"]
    rPr = r_el.find(f"{{{ns_w}}}rPr")
    if rPr is None:
        return None
    hl_el = rPr.find(f"{{{ns_w}}}highlight")
    if hl_el is None:
        return None
    return hl_el.get(f"{{{ns_w}}}val")


def _norm_highlight(v: Optional[str]) -> Optional[str]:
    """Collapse invisible highlight values to None so they never count as contrast."""
    if v is None:
        return None
    v = v.lower()
    return None if v in _INVISIBLE_HIGHLIGHTS else v


def _dominant_highlight_over_runs(r_els) -> Optional[str]:
    """Most common highlight across runs, weighted by visible char count.
    None when unhighlighted ink dominates — the common case."""
    counts: Counter = Counter()
    for r in r_els:
        n = _run_visible_len(r)
        if n:
            counts[_norm_highlight(_run_highlight(r))] += n
    if not counts:
        return None
    return counts.most_common(1)[0][0]


def _body_dominant_highlight(body) -> Optional[str]:
    ns_w = _NS["w"]
    runs: List = []
    for child in body:
        if _tag(child) == "p":
            runs.extend(child.iter(f"{{{ns_w}}}r"))
    return _dominant_highlight_over_runs(runs)


def _table_dominant_highlight(table_el) -> Optional[str]:
    return _dominant_highlight_over_runs(table_el.iter(f"{{{_NS['w']}}}r"))


def _dominant_color_over_runs(r_els) -> Optional[str]:
    """Most common color across runs, weighted by visible char count.

    Returns None when black-ish (default) ink dominates — the common case for prose.
    """
    counts: Counter = Counter()
    for r in r_els:
        n = _run_visible_len(r)
        if n:
            counts[_norm_color(_run_color(r))] += n
    if not counts:
        return None
    return counts.most_common(1)[0][0]


def _body_dominant(body) -> Optional[str]:
    """Dominant text color of the body's top-level paragraphs (prose, excluding tables)."""
    ns_w = _NS["w"]
    runs: List = []
    for child in body:
        if _tag(child) == "p":
            runs.extend(child.iter(f"{{{ns_w}}}r"))
    return _dominant_color_over_runs(runs)


def _table_dominant(table_el) -> Optional[str]:
    """Dominant text color across all runs in a table element."""
    return _dominant_color_over_runs(table_el.iter(f"{{{_NS['w']}}}r"))


# =============================================================================
# RENDER STATS
# =============================================================================

@dataclass
class RenderStats:
    total_paragraphs: int = 0
    total_tables: int = 0
    nested_tables: int = 0
    images_found: int = 0
    textboxes_found: int = 0
    merged_cells: int = 0
    strikethrough_runs: int = 0
    color_spans: int = 0
    highlight_spans: int = 0
    empty_skipped: int = 0
    rendered_chars: int = 0

    def log(self) -> None:
        logger.info(
            f"[RENDER] {self.rendered_chars:,} chars | "
            f"{self.total_paragraphs} paragraphs | "
            f"{self.total_tables} tables ({self.nested_tables} nested) | "
            f"{self.images_found} images | "
            f"{self.textboxes_found} textboxes | "
            f"{self.merged_cells} merged cells | "
            f"{self.strikethrough_runs} strikethrough spans | "
            f"{self.color_spans} color spans | "
            f"{self.highlight_spans} highlight spans"
        )


# =============================================================================
# PARAGRAPH HELPERS
# =============================================================================

def _para_text(p_el, baseline: Optional[str] = None,
               hl_baseline: Optional[str] = None) -> str:
    """Extract text from a <w:p> element, preserving strikethrough as ~~text~~,
    contrast color as [[color:RRGGBB]]text[[/color]], and contrast highlight as
    [[hl:name]]text[[/hl]].

    Iterates <w:r> (run) elements to detect per-run formatting. Strikethrough
    (<w:strike/> or <w:dStrike/>) renders as ~~strikethrough~~; a run colored
    differently from `baseline` (the enclosing scope's dominant color) renders as
    [[color:RRGGBB]]...[[/color]]; a run highlighted differently from `hl_baseline`
    (the scope's dominant highlight; white/none never count) renders as
    [[hl:name]]...[[/hl]] so the LLM can see teacher annotations on either physical
    channel. Black-ish colors and baseline-matching runs are left unmarked.
    """
    parts: List[str] = []
    ns_w = _NS["w"]
    base = _norm_color(baseline)
    hl_base = _norm_highlight(hl_baseline)

    # Collect per-run (is_strike, contrast_color, contrast_hl, text), then merge
    # consecutive runs with the same key so a red solution split across many runs
    # becomes ONE span ([[color:EE0000]]…[[/color]]) instead of one per character.
    segments: List[Tuple[bool, Optional[str], Optional[str], str]] = []

    for child in p_el:
        tag = _tag(child)

        if tag == "r":
            is_strike = False
            rPr = child.find(f"{{{ns_w}}}rPr")
            if rPr is not None:
                if rPr.find(f"{{{ns_w}}}strike") is not None:
                    is_strike = True
                if rPr.find(f"{{{ns_w}}}dstrike") is not None:
                    is_strike = True

            run_text_parts: List[str] = []
            for el in child:
                el_tag = _tag(el)
                if el_tag == "t" and el.text:
                    run_text_parts.append(el.text)
                elif el_tag == "tab":
                    run_text_parts.append("\t")
                elif el_tag == "br":
                    run_text_parts.append("\n")

            run_text = "".join(run_text_parts)
            if run_text:
                color = _norm_color(_run_color(child))
                contrast = color if (color is not None and color != base
                                     and _is_annotation_color(color)) else None
                # Whitespace-only runs carry no annotation information on the hl
                # channel — marking them yields junk spans like "[[hl:yellow]] [[/hl]]"
                # (color/strike whitespace handling is left untouched: pre-B1 behavior).
                hl = _norm_highlight(_run_highlight(child))
                contrast_hl = hl if (hl is not None and hl != hl_base
                                     and run_text.strip()) else None
                segments.append((is_strike, contrast, contrast_hl, run_text))

        elif tag == "hyperlink":
            for r in child:
                if _tag(r) == "r":
                    htext = "".join(el.text for el in r
                                    if _tag(el) == "t" and el.text)
                    if htext:
                        segments.append((False, None, None, htext))

    # Merge consecutive segments sharing (strike, contrast, hl) and emit markup.
    merged: List[Tuple[Tuple[bool, Optional[str], Optional[str]], List[str]]] = []
    for strike, contrast, hl, text in segments:
        key = (strike, contrast, hl)
        if merged and merged[-1][0] == key:
            merged[-1][1].append(text)
        else:
            merged.append((key, [text]))

    for (strike, contrast, hl), texts in merged:
        s = "".join(texts)
        if strike:
            s = s.strip()
            if hl is not None:
                s = f"[[hl:{hl}]]{s}[[/hl]]"
            if contrast is not None:
                s = f"[[color:{contrast}]]{s}[[/color]]"
            s = f"~~{s}~~"
        else:
            if hl is not None:
                s = f"[[hl:{hl}]]{s}[[/hl]]"
            if contrast is not None:
                s = f"[[color:{contrast}]]{s}[[/color]]"
        parts.append(s)

    return "".join(parts).strip()


def _para_images(p_el) -> List[str]:
    """Detect images in a paragraph. Returns image name list."""
    imgs: List[str] = []
    for drawing in p_el.iter("{%s}drawing" % _NS["w"]):
        for dp in drawing.iter("{%s}docPr" % _NS["wp"]):
            imgs.append(dp.get("descr") or dp.get("name") or "image")
            break
        else:
            if list(drawing.iter("{%s}blip" % _NS["a"])):
                imgs.append("image")
    for alt in p_el.iter("{%s}AlternateContent" % _NS["mc"]):
        if list(alt.iter("{%s}blip" % _NS["a"])):
            imgs.append("image")
    return imgs


def _para_textboxes(p_el) -> List[str]:
    """Extract textbox content from drawing shapes."""
    results: List[str] = []
    for txbx in p_el.iter("{%s}txbx" % _NS["wps"]):
        parts: List[str] = []
        for el in txbx.iter():
            tag = _tag(el)
            if tag == "t" and el.text:
                parts.append(el.text)
        text = " ".join(parts).strip()
        if text:
            results.append(text)
    return results


def _render_para(p_el, stats: RenderStats, baseline: Optional[str] = None,
                 hl_baseline: Optional[str] = None) -> List[str]:
    """Render a paragraph to lines."""
    lines: List[str] = []

    for img in _para_images(p_el):
        stats.images_found += 1
        lines.append(f"[IMAGE: {img}]")

    for tb in _para_textboxes(p_el):
        stats.textboxes_found += 1
        lines.append(f"[TEXTBOX]\n{tb}\n[/TEXTBOX]")

    text = _para_text(p_el, baseline, hl_baseline)
    if text:
        stats.total_paragraphs += 1
        lines.append(text)
    elif not lines:
        stats.empty_skipped += 1

    return lines


# =============================================================================
# CELL CONTENT ANALYSIS
# =============================================================================

def _analyze_cell(cell, baseline: Optional[str] = None,
                  hl_baseline: Optional[str] = None) -> Tuple[str, List]:
    """Analyze a cell's content: separate text from nested tables.

    Returns:
        (cell_text, nested_table_elements)
        - cell_text: all paragraph text in the cell (no nested table content)
        - nested_table_elements: list of <w:tbl> elements found in the cell
    """
    text_parts: List[str] = []
    nested_tbls = []

    for child in cell._element:
        tag = _tag(child)
        if tag == "p":
            t = _para_text(child, baseline, hl_baseline)
            if t:
                text_parts.append(t)
            # Also capture image placeholders from cell paragraphs
            for img in _para_images(child):
                text_parts.append(f"[IMAGE: {img}]")
        elif tag == "tbl":
            nested_tbls.append(child)
        # tcPr and others are skipped

    return " ".join(text_parts), nested_tbls


def _clean(text: str) -> str:
    """Clean text for markdown cell rendering."""
    cleaned = " ".join(text.split())
    return cleaned.replace("|", "\\|")


# =============================================================================
# TABLE RENDERING
# =============================================================================

def _render_table(
    table: DocxTableObj,
    index: Optional[int],
    stats: RenderStats,
    indent: str = "",
) -> List[str]:
    """Render a table to markdown lines. Nested tables become indented blocks.
    
    Returns list of lines (not joined — caller decides spacing).
    """
    rows = table.rows
    if not rows:
        label = f"TABLE {index}" if index else "NESTED TABLE"
        return [f"{indent}[{label}: empty]"]

    stats.total_tables += 1
    num_rows = len(rows)
    num_cols = max(len(row.cells) for row in rows)

    # Contrast baselines for THIS table (nested tables compute their own below).
    # A uniformly-colored/highlighted table yields a baseline equal to its runs -> no marks.
    table_baseline = _table_dominant(table._element)
    table_hl_baseline = _table_dominant_highlight(table._element)

    lines: List[str] = []

    # Label
    if index is not None:
        lines.append(f"{indent}[TABLE {index}: {num_rows}x{num_cols}]")
    else:
        lines.append(f"{indent}[NESTED TABLE: {num_rows}x{num_cols}]")

    # Process header row
    header_cells, header_nested = _process_row(rows[0], stats, table_baseline, table_hl_baseline)
    # Trim trailing empty cells from merges
    while header_cells and header_cells[-1] == "":
        header_cells.pop()
    if not header_cells:
        header_cells = ["(empty)"]

    lines.append(f"{indent}| " + " | ".join(header_cells) + " |")
    lines.append(f"{indent}|" + "|".join(["---"] * len(header_cells)) + "|")

    # Render nested tables from header (rare but possible)
    for nt_el in header_nested:
        lines.extend(_render_nested_table_element(table, nt_el, stats, indent))

    # Data rows
    max_rows = min(num_rows, 60)
    for row in rows[1:max_rows]:
        cells, row_nested = _process_row(row, stats, table_baseline, table_hl_baseline)
        # Pad/trim to header width
        while len(cells) < len(header_cells):
            cells.append("")
        cells = cells[:len(header_cells)]

        lines.append(f"{indent}| " + " | ".join(cells) + " |")

        # Render any nested tables found in this row's cells
        for nt_el in row_nested:
            lines.extend(_render_nested_table_element(table, nt_el, stats, indent))

    if num_rows > max_rows:
        lines.append(f"{indent}... ({num_rows - max_rows} more rows)")

    return lines


def _process_row(row, stats: RenderStats, baseline: Optional[str] = None,
                 hl_baseline: Optional[str] = None) -> Tuple[List[str], List]:
    """Process a table row: extract cell texts and collect nested table elements.

    Returns:
        (cell_texts, all_nested_table_elements)
    """
    cell_texts: List[str] = []
    all_nested: List = []
    seen_ids: set = set()

    for cell in row.cells:
        cell_id = id(cell._element)
        if cell_id in seen_ids:
            # Merged cell — same element appears multiple times
            cell_texts.append("")
            stats.merged_cells += 1
            continue
        seen_ids.add(cell_id)

        text, nested = _analyze_cell(cell, baseline, hl_baseline)
        cell_texts.append(_clean(text))
        all_nested.extend(nested)

    return cell_texts, all_nested


def _render_nested_table_element(
    parent_table: DocxTableObj,
    tbl_element,
    stats: RenderStats,
    parent_indent: str,
) -> List[str]:
    """Render a nested <w:tbl> element as an indented block."""
    stats.nested_tables += 1
    child_indent = parent_indent + "  "

    # Convert element to python-docx Table
    try:
        nested_table = DocxTableObj(tbl_element, parent_table._element)
    except Exception:
        # Fallback: try to find it in parent cells
        for row in parent_table.rows:
            for cell in row.cells:
                for ct in cell.tables:
                    if ct._element is tbl_element:
                        nested_table = ct
                        break
                else:
                    continue
                break
            else:
                continue
            break
        else:
            return [f"{child_indent}[NESTED TABLE: render failed]"]

    return _render_table(nested_table, index=None, stats=stats, indent=child_indent)


# =============================================================================
# MAIN ENTRY POINT
# =============================================================================

def render_docx_to_markdown(file_bytes: bytes) -> str:
    """Convert DOCX bytes to LLM-ready markdown.

    All elements rendered in document order. Nested tables appear as
    indented blocks below their parent row for clear LLM readability.

    Returns:
        Rendered markdown string.
    Raises:
        ValueError: If DOCX cannot be parsed.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}") from e

    stats = RenderStats()
    body = doc.element.body
    tables_by_el = {t._element: t for t in doc.tables}

    # Contrast baselines for body prose (computed once; tables use their own).
    body_baseline = _body_dominant(body)
    body_hl_baseline = _body_dominant_highlight(body)

    output: List[str] = []
    table_index = 0

    for child in body:
        tag = _tag(child)

        if tag == "p":
            for line in _render_para(child, stats, body_baseline, body_hl_baseline):
                output.append(line)
                output.append("")

        elif tag == "tbl":
            table_index += 1
            tbl_obj = tables_by_el.get(child)
            if tbl_obj:
                table_lines = _render_table(tbl_obj, table_index, stats)
                output.extend(table_lines)
                output.append("")

        elif tag == "sdt":
            # Structured document tags — extract inner paragraphs
            for inner in child.iter():
                if _tag(inner) == "p":
                    for line in _render_para(inner, stats, body_baseline, body_hl_baseline):
                        output.append(line)
                        output.append("")

    result = "\n".join(output)
    stats.rendered_chars = len(result)
    # Span counts from the result itself (the strikethrough_runs counter predates the
    # run-merging and was never incremented; counting emitted spans is the honest metric).
    stats.strikethrough_runs = result.count("~~") // 2
    stats.color_spans = result.count("[[color:")
    stats.highlight_spans = result.count("[[hl:")

    # Observability
    stats.log()
    _log_preview(result)

    # Debug artifact — best-effort ONLY. A relative-path write fails when CWD isn't
    # backend/ or the filesystem is read-only (Cloud Run); a debug file must never
    # fail a production render.
    try:
        with open("app/services/docx_v3/rendered_output.md", "w", encoding="utf-8") as f:
            f.write(result)
    except OSError as e:
        logger.debug(f"[RENDER] debug artifact write skipped: {e}")

    return result


def audit_annotation_channels(file_bytes: bytes, rendered: str) -> List[str]:
    """Formatting-annotation loss guard (B1b). PURE; no side effects.

    Text-based render-loss attribution is blind to FORMATTING loss: a dropped
    [[color]]/[[hl]]/~~ mark loses teacher ink while every character still renders
    (exactly how the highlight channel stayed invisible until foundations q1.ג).
    This audit counts annotation-bearing runs on the DOCX side — applying the SAME
    per-scope contrast rules as the renderer, so legitimately-suppressed uniform
    scopes (csharp's all-red table) do not false-positive — and warns when a channel
    has docx ink but ZERO rendered marks. Run counts vs merged span counts differ by
    design; the guard is presence-vs-zero per channel, not numeric equality.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        return []  # unparseable docx fails elsewhere; the audit never masks that error

    body = doc.element.body
    ns_w = _NS["w"]
    body_base = _norm_color(_body_dominant(body))
    body_hl_base = _norm_highlight(_body_dominant_highlight(body))

    # Per-table baselines; a run's scope = nearest ancestor <w:tbl>, else body.
    tbl_bases = {}
    for tbl in body.iter(f"{{{ns_w}}}tbl"):
        tbl_bases[tbl] = (_norm_color(_table_dominant(tbl)),
                          _norm_highlight(_table_dominant_highlight(tbl)))

    def _scope_bases(r_el):
        for anc in r_el.iterancestors():
            if anc in tbl_bases:
                return tbl_bases[anc]
        return body_base, body_hl_base

    docx_counts = {"strike": 0, "color": 0, "highlight": 0}
    for r in body.iter(f"{{{ns_w}}}r"):
        if not _run_visible_len(r):
            continue
        rPr = r.find(f"{{{ns_w}}}rPr")
        if rPr is not None and (rPr.find(f"{{{ns_w}}}strike") is not None
                                or rPr.find(f"{{{ns_w}}}dstrike") is not None):
            docx_counts["strike"] += 1
        base, hl_base = _scope_bases(r)
        color = _norm_color(_run_color(r))
        if color is not None and color != base and _is_annotation_color(color):
            docx_counts["color"] += 1
        hl = _norm_highlight(_run_highlight(r))
        if hl is not None and hl != hl_base:
            docx_counts["highlight"] += 1

    rendered_counts = {"strike": rendered.count("~~") // 2,
                       "color": rendered.count("[[color:"),
                       "highlight": rendered.count("[[hl:")}
    warnings = []
    for ch, n in docx_counts.items():
        if n > 0 and rendered_counts[ch] == 0:
            warnings.append(
                f"render annotation loss: docx has {n} '{ch}'-annotated runs "
                f"but the render emitted 0 {ch} marks — teacher ink on the {ch} "
                f"channel never reached the LLM (parser_render bug)")
    return warnings


def _log_preview(rendered: str) -> None:
    """Log what the LLM will see: questions, tables, images."""
    questions = []
    tables = []
    images = []

    for line in rendered.split("\n"):
        s = line.strip()
        if s.startswith("שאלה ") and any(c.isdigit() for c in s):
            questions.append(s[:80])
        elif s.startswith("[TABLE ") or s.startswith("[NESTED TABLE"):
            tables.append(s[:80])
        elif s.startswith("[IMAGE:"):
            images.append(s)

    logger.info("[RENDER] Structure:")
    for q in questions:
        logger.info(f"  Q: {q}")
    for t in tables:
        logger.info(f"  T: {t}")
    if images:
        logger.info(f"  Images: {len(images)}")